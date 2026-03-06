#!/usr/bin/env python3
"""Document reader with format auto-detection.

Supports: .docx, .xlsx, .csv, .tsv, .pdf, .json, .yaml, .md
Reads file and outputs text content with truncation.
---
description: Read document files (.docx, .xlsx, .pdf, .csv, .json, .yaml)
databases: []
read_only: true
---
"""

import json
import sys
import io
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))

from lib.io import output


def read_document_impl(path=None, max_chars=50000, sheet=None, pages=None,
                       driver=None, **kwargs):
    """Read a document file and return its text content.

    Args:
        path: Absolute path to the file
        max_chars: Maximum characters to return (default 50000)
        sheet: For .xlsx - sheet name to read (default: active sheet)
        pages: For .pdf - page range like '1-5' or '3' (default: all)
        driver: Ignored (accepted for dispatch compatibility)

    Returns:
        str: Document text content or error message
    """
    if not path:
        return "ERROR: No path provided"

    filepath = Path(path)
    if not filepath.exists():
        return f"ERROR: File not found: {path}"

    ext = filepath.suffix.lower()
    output_parts = []

    try:
        if ext == '.docx':
            from docx import Document
            doc = Document(str(filepath))
            for para in doc.paragraphs:
                if para.text.strip():
                    output_parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    output_parts.append(" | ".join(cells))

        elif ext in ('.xlsx', '.xls'):
            import openpyxl
            wb = openpyxl.load_workbook(str(filepath), read_only=True, data_only=True)
            sheet_name = sheet or wb.active.title
            if sheet_name not in wb.sheetnames:
                output_parts.append(f"Available sheets: {', '.join(wb.sheetnames)}")
                sheet_name = wb.active.title
            ws = wb[sheet_name]
            output_parts.append(f"[Sheet: {sheet_name}]")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                output_parts.append("\t".join(cells))
            wb.close()

        elif ext in ('.csv', '.tsv'):
            import csv
            delimiter = '\t' if ext == '.tsv' else ','
            with open(str(filepath), 'r', encoding='utf-8', errors='replace') as cf:
                reader = csv.reader(cf, delimiter=delimiter)
                for row in reader:
                    output_parts.append("\t".join(row))

        elif ext == '.pdf':
            import pdfplumber
            with pdfplumber.open(str(filepath)) as pdf:
                page_range = None
                if pages:
                    parts = str(pages).split("-")
                    if len(parts) == 2:
                        page_range = range(int(parts[0]) - 1, int(parts[1]))
                    else:
                        page_range = range(int(parts[0]) - 1, int(parts[0]))
                else:
                    page_range = range(len(pdf.pages))

                output_parts.append(f"[PDF: {len(pdf.pages)} pages total]")
                for i in page_range:
                    if i < len(pdf.pages):
                        text = pdf.pages[i].extract_text()
                        if text:
                            output_parts.append(f"--- Page {i+1} ---")
                            output_parts.append(text)

        elif ext == '.json':
            with open(str(filepath), 'r', encoding='utf-8') as jf:
                data = json.load(jf)
                output_parts.append(json.dumps(data, indent=2, ensure_ascii=False))

        elif ext in ('.yaml', '.yml'):
            import yaml
            with open(str(filepath), 'r', encoding='utf-8') as yf:
                data = yaml.safe_load(yf)
                output_parts.append(json.dumps(data, indent=2, ensure_ascii=False, default=str))

        elif ext == '.md':
            import frontmatter
            post = frontmatter.load(str(filepath))
            if post.metadata:
                output_parts.append("[FRONTMATTER]")
                output_parts.append(json.dumps(dict(post.metadata), indent=2, ensure_ascii=False, default=str))
                output_parts.append("[CONTENT]")
            output_parts.append(post.content)

        else:
            text = filepath.read_text(encoding='utf-8', errors='replace')
            output_parts.append(text)

    except Exception as e:
        return f"ERROR reading {ext} file: {e}"

    result = "\n".join(output_parts)
    if len(result) > max_chars:
        result = result[:max_chars] + f"\n\n[TRUNCATED at {max_chars} chars - {len(result)} total]"

    return result


def main():
    """Subprocess entry point: read params from JSON file."""
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    if len(sys.argv) < 2:
        print("ERROR: Missing params file path", file=sys.stderr)
        sys.exit(1)

    try:
        with open(sys.argv[1], 'r', encoding='utf-8') as f:
            p = json.load(f)
    except Exception as e:
        print(f"ERROR: Failed to load params file: {e}", file=sys.stderr)
        sys.exit(1)

    result = read_document_impl(**p)
    print(result)


if __name__ == "__main__":
    main()
